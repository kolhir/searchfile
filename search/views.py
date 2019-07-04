from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.template import RequestContext
from django.http import HttpResponseRedirect, HttpResponse, Http404
from django.shortcuts import render_to_response
from .forms import UploadFileForm
from .models import Files, save_file
import docx
import os
from django.conf import settings
from searchfile.settings import MEDIA_ROOT


@csrf_exempt
def index(request):
    #вьюшка главной страницы, с добавлением файла и полем для поиска
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            #вызываем метод для сохранения файла из .models
            save_file(request.POST.get("title"), request.FILES['file'])
            return HttpResponseRedirect('/')
    else:
        form = UploadFileForm()
    return render_to_response('index.html', {"form":form})


@csrf_exempt
def search(request):
    #обработка поискового запроса
    if request.method == 'POST':
        search = request.POST.get("search")
        docs = Files.objects.all()
        res_dict = {} #словарь для хранения совпадений
        #первый цикл, проходит по всем файлам
        for doc in docs:
            res_dict.update({doc:[]})
            print(doc.file.url)
            file = docx.Document(doc.file.path)
            #второй цикл, ищем в каждом параграфе совпадения
            for i, p in enumerate(file.paragraphs):
                #если находим совпадение, добавляем в словарь параграф +-1
                if search in p.text:
                    s = file.paragraphs[i-1].text + p.text + file.paragraphs[i-1].text
                    res_dict[doc].append(s)
        return render_to_response('result.html', {"results":res_dict})

def download(request, id, name):
    #скачиванием файла
    file =  Files.objects.get(id = id)
    path = file.file.path
    file_path = os.path.join(MEDIA_ROOT, path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
            return response
    raise Http404